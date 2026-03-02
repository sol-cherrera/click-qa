"""
src/exporters/docx_exporter.py — Exporta pasos a DOCX con python-docx.
"""
import os
import tempfile
from datetime import datetime
from typing import List, Dict, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def export_docx(steps: List[Dict[str, Any]], filepath: str) -> None:
    """
    Genera un documento Word con todos los pasos.
    Incluye: portada, descripcion del paso, timestamp, imagen embebida.
    """
    doc = Document()

    # ── Estilos globales ──────────────────────────────────────────
    _configure_styles(doc)

    # ── Portada ───────────────────────────────────────────────────
    _add_cover(doc, steps)

    # ── Pasos ─────────────────────────────────────────────────────
    tmp_files = []
    try:
        for step in steps:
            doc.add_page_break()

            # Encabezado del paso
            hdr = doc.add_paragraph()
            hdr.paragraph_format.space_before = Pt(6)
            hdr.paragraph_format.space_after  = Pt(4)
            hdr_run = hdr.add_run(f"Paso {step['number']}")
            hdr_run.bold      = True
            hdr_run.font.size = Pt(14)
            hdr_run.font.color.rgb = RGBColor(124, 58, 237)

            action_run = hdr.add_run(f"  ·  {step.get('action','Accion')}")
            action_run.font.size = Pt(10)
            action_run.font.color.rgb = RGBColor(167, 139, 250)

            # Descripcion
            desc_p = doc.add_paragraph()
            desc_p.paragraph_format.space_before = Pt(2)
            desc_p.paragraph_format.space_after  = Pt(4)
            desc_run = desc_p.add_run(step.get("description", ""))
            desc_run.bold      = True
            desc_run.font.size = Pt(12)

            # Timestamp
            if step.get("timestamp"):
                try:
                    ts = datetime.fromisoformat(step["timestamp"])
                    ts_str = ts.strftime("%H:%M:%S — %d/%m/%Y")
                except Exception:
                    ts_str = step["timestamp"]
                ts_p = doc.add_paragraph()
                ts_p.paragraph_format.space_after = Pt(6)
                ts_run = ts_p.add_run(ts_str)
                ts_run.font.size   = Pt(9)
                ts_run.font.italic = True
                ts_run.font.color.rgb = RGBColor(130, 130, 150)

            # Imagen
            if step.get("screenshot_bytes"):
                tmp = _save_temp_image(step["screenshot_bytes"])
                if tmp:
                    tmp_files.append(tmp)
                    try:
                        img_p = doc.add_paragraph()
                        img_p.paragraph_format.space_before = Pt(4)
                        run = img_p.add_run()
                        run.add_picture(tmp, width=Inches(6.2))
                    except Exception as e:
                        err_p = doc.add_paragraph(f"[Error al cargar imagen: {e}]")
                        err_p.runs[0].font.color.rgb = RGBColor(200, 80, 80)

    finally:
        for p in tmp_files:
            try:
                os.unlink(p)
            except Exception:
                pass

    doc.save(filepath)


# ─── Helpers ──────────────────────────────────────────────────────
def _add_cover(doc: Document, steps: List[Dict[str, Any]]):
    """Pagina de portada."""
    doc.add_paragraph()  # espacio superior

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    t_run = title.add_run("Plan de Prueba Manual")
    t_run.bold = True
    t_run.font.size = Pt(24)
    t_run.font.color.rgb = RGBColor(124, 58, 237)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(8)
    s_run = sub.add_run(
        f"{len(steps)} paso{'s' if len(steps)!=1 else ''} documentado{'s' if len(steps)!=1 else ''}"
    )
    s_run.font.size = Pt(13)
    s_run.font.color.rgb = RGBColor(136, 136, 153)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(4)
    d_run = date_p.add_run(f"Generado el {datetime.now().strftime('%d/%m/%Y — %H:%M')}")
    d_run.font.size = Pt(10)
    d_run.font.color.rgb = RGBColor(100, 100, 120)


def _configure_styles(doc: Document):
    """Configurar fuente y margenes del documento."""
    section = doc.sections[0]
    section.page_width  = Cm(21)    # A4
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _save_temp_image(screenshot_bytes: bytes) -> str | None:
    try:
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp.write(screenshot_bytes)
        tmp.close()
        return tmp.name
    except Exception:
        return None
